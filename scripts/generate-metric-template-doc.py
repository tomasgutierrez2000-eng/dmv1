#!/usr/bin/env python3
"""Generate the Dynamic Metric Generation System template Word document."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import os

doc = Document()

# ── Styles ──────────────────────────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

for level in range(1, 5):
    h = doc.styles[f'Heading {level}']
    h.font.color.rgb = RGBColor(0x1a, 0x36, 0x5d)  # Dark navy

def add_table_with_style(doc, rows, cols, data=None, header=True, col_widths=None):
    """Add a formatted table."""
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    if data:
        for r_idx, row_data in enumerate(data):
            for c_idx, cell_val in enumerate(row_data):
                cell = table.cell(r_idx, c_idx)
                cell.text = str(cell_val)
                if r_idx == 0 and header:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.bold = True
                            run.font.size = Pt(10)
                else:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.font.size = Pt(10)
    return table

def add_code_block(doc, text):
    """Add a monospace code block."""
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

def add_note(doc, text, label="NOTE"):
    """Add a highlighted note."""
    p = doc.add_paragraph()
    run = p.add_run(f'{label}: ')
    run.bold = True
    run.font.color.rgb = RGBColor(0xc0, 0x39, 0x2b)
    run2 = p.add_run(text)
    run2.font.size = Pt(10)
    run2.font.italic = True


# ════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_heading('Dynamic Metric Generation System', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_paragraph('Template, Sample & End-to-End Instructions')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in subtitle.runs:
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run('GSIB Banking Data Model Visualizer\n').font.size = Pt(12)
meta.add_run('Version 1.0 — March 2026\n').font.size = Pt(11)
meta.add_run('\nFor use with: L1/L2/L3 Three-Layer Data Architecture').font.size = Pt(10)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ════════════════════════════════════════════════════════════════════════════
doc.add_heading('Table of Contents', level=1)
toc_items = [
    ('1', 'System Overview'),
    ('2', 'Architecture: The Three-Layer Data Model'),
    ('3', 'Rollup Hierarchy'),
    ('4', 'Metric Template — Blank'),
    ('  4.1', 'Section A: Metric Identity'),
    ('  4.2', 'Section B: Classification & Metadata'),
    ('  4.3', 'Section C: Ingredient Fields (Source Data)'),
    ('  4.4', 'Section D: Level Definitions (Rollup Logic)'),
    ('  4.5', 'Section E: Executable Specification'),
    ('  4.6', 'Section F: Demo Data (Optional)'),
    ('5', 'Sample Filled Template — Debt Service Coverage Ratio (DSCR)'),
    ('6', 'End-to-End Instructions'),
    ('  6.1', 'Step 1: Define the Business Concept'),
    ('  6.2', 'Step 2: Identify Ingredient Fields'),
    ('  6.3', 'Step 3: Design Rollup Logic'),
    ('  6.4', 'Step 4: Write Executable SQL'),
    ('  6.5', 'Step 5: Validate & Review'),
    ('  6.6', 'Step 6: Register in the System'),
    ('7', 'Rollup Review Checklist for GSIB'),
    ('8', 'Domain Reference'),
    ('9', 'Field Naming Convention Reference'),
    ('A', 'Appendix: Sourcing Type Decision Tree'),
]
for num, title_text in toc_items:
    p = doc.add_paragraph()
    p.add_run(f'{num}  ').bold = True
    p.add_run(title_text)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 1. SYSTEM OVERVIEW
# ════════════════════════════════════════════════════════════════════════════
doc.add_heading('1. System Overview', level=1)
doc.add_paragraph(
    'The Dynamic Metric Generation System enables business analysts and risk managers '
    'to define, calculate, and visualize metrics across a GSIB banking data model. '
    'Each metric is defined once as a business concept (CatalogueItem) and optionally '
    'linked to an executable specification (L3Metric) that runs SQL against the data warehouse.'
)
doc.add_paragraph(
    'This document provides: (1) a blank template for defining new metrics, '
    '(2) a fully worked example using DSCR, and (3) step-by-step instructions '
    'for filling out the template and reviewing rollups end-to-end.'
)

doc.add_heading('Key Concepts', level=2)
concepts = [
    ['Concept', 'Description'],
    ['CatalogueItem', 'Business definition of a metric — identity, formula, rollup logic, source fields'],
    ['L3Metric', 'Executable specification — SQL formulas per dimension, display formats, lineage'],
    ['Level Definition', 'How the metric computes at each rollup level (facility → counterparty → desk → portfolio → LOB)'],
    ['Ingredient Field', 'An atomic source field from an L1 or L2 table used in the calculation'],
    ['Sourcing Type', 'How data arrives at each level: Raw, Calc, Agg, or Avg'],
    ['Domain', 'Business area classification (Credit Risk, Exposure & Limits, Financial Performance, etc.)'],
]
add_table_with_style(doc, len(concepts), 2, concepts)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 2. ARCHITECTURE
# ════════════════════════════════════════════════════════════════════════════
doc.add_heading('2. Architecture: The Three-Layer Data Model', level=1)

layers = [
    ['Layer', 'Purpose', 'Examples', 'Table Count'],
    ['L1 — Reference Data', 'Dimensions, masters, lookups, hierarchies, configuration. Rarely changes.', 'counterparty, facility_master, currency_dim, metric_threshold', '82'],
    ['L2 — Atomic Data', 'Raw source-system snapshots and events. Point-in-time observations, not computed.', 'facility_exposure_snapshot, credit_event, position', '25'],
    ['L3 — Derived Data', 'Anything calculated, aggregated, or computed from L1+L2.', 'exposure_metric_cube, facility_financial_calc, stress_test_result', '54'],
]
add_table_with_style(doc, len(layers), 4, layers)

doc.add_paragraph()
add_note(doc, 'Data flows forward only: L1 → L2 → L3. Never backwards. If a field is computed from other fields (ratios, aggregations), it belongs in L3.', 'RULE')

doc.add_heading('Calculated Overlay Pattern', level=2)
doc.add_paragraph(
    'When an L2 table has a mix of raw and derived fields, split the derived fields into '
    'a new L3 table at the same grain (same PK) with a FK back to the L2 source.'
)
doc.add_paragraph('Example: l2.facility_financial_snapshot (raw inputs) → l3.facility_financial_calc (DSCR, LTV, net income)')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 3. ROLLUP HIERARCHY
# ════════════════════════════════════════════════════════════════════════════
doc.add_heading('3. Rollup Hierarchy', level=1)
doc.add_paragraph(
    'Every metric must define how it computes at each level of the rollup hierarchy. '
    'The hierarchy flows from the most granular (facility) to the most aggregated (business segment).'
)

hierarchy = [
    ['Level', 'Grain', 'Key Table', 'Description'],
    ['Facility', 'Individual credit facility', 'l1.facility_master', 'Most granular level — one row per facility'],
    ['Counterparty', 'Legal entity / borrower', 'l1.counterparty', 'Aggregates all facilities for a borrower'],
    ['Desk (L3)', 'Trading / origination desk', 'l1.enterprise_business_taxonomy (tree_level=L3)', 'Organizational unit managing the exposure'],
    ['Portfolio (L2)', 'Portfolio grouping', 'l1.enterprise_business_taxonomy (tree_level=L2)', 'Mid-level organizational aggregation'],
    ['Business Segment (L1)', 'Line of business', 'l1.enterprise_business_taxonomy (tree_level=L1)', 'Highest-level organizational aggregation'],
]
add_table_with_style(doc, len(hierarchy), 4, hierarchy)

doc.add_paragraph()
doc.add_paragraph('Rollup direction: Facility → Counterparty → Desk (L3) → Portfolio (L2) → Business Segment (L1)')

doc.add_heading('Sourcing Types', level=2)
sourcing = [
    ['Type', 'Code', 'When to Use', 'Example'],
    ['Raw', 'Raw', 'Field is read directly from L1/L2 with no computation', 'committed_facility_amt from facility_master'],
    ['Calculated', 'Calc', 'Formula applied at this level using atomic fields', 'DSCR = net_operating_income / total_debt_service at facility level'],
    ['Aggregated', 'Agg', 'SUM, COUNT, MAX, MIN of lower-level values', 'Total Exposure = SUM(facility exposures) at desk level'],
    ['Weighted Avg', 'Avg', 'Exposure-weighted or EAD-weighted average', 'Portfolio PD = SUM(PD * EAD) / SUM(EAD)'],
]
add_table_with_style(doc, len(sourcing), 4, sourcing)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 4. METRIC TEMPLATE — BLANK
# ════════════════════════════════════════════════════════════════════════════
doc.add_heading('4. Metric Template — Blank', level=1)
doc.add_paragraph(
    'Complete each section below when defining a new metric. '
    'Fields marked with (*) are required. All others are optional but recommended.'
)

# Section A
doc.add_heading('4.1  Section A: Metric Identity', level=2)
identity_fields = [
    ['Field', 'Value', 'Instructions'],
    ['item_id *', '', 'Unique identifier. Use abbreviated form: "DSCR", "LTV", "PD", or "MET-XXX" for new metrics'],
    ['item_name *', '', 'Full business name: "Debt Service Coverage Ratio"'],
    ['abbreviation *', '', 'Short display label: "DSCR", "LTV"'],
    ['kind *', '', 'METRIC (calculated) or DATA_ELEMENT (raw field)'],
    ['definition *', '', 'Complete business definition (2-4 sentences) including regulatory context'],
    ['generic_formula *', '', 'High-level formula in plain English: "Net Operating Income / Total Debt Service"'],
    ['insight *', '', 'One-sentence executive summary for dashboards'],
    ['status *', '', 'ACTIVE, DRAFT, or DEPRECATED'],
]
add_table_with_style(doc, len(identity_fields), 3, identity_fields)

# Section B
doc.add_heading('4.2  Section B: Classification & Metadata', level=2)
class_fields = [
    ['Field', 'Value', 'Allowed Values'],
    ['data_type *', '', 'Decimal, Integer, String, Boolean, Date'],
    ['unit_type *', '', 'RATIO, PERCENTAGE, CURRENCY, COUNT, RATE, ORDINAL, DAYS, INDEX'],
    ['direction *', '', 'HIGHER_BETTER, LOWER_BETTER, NEUTRAL'],
    ['metric_class *', '', 'SOURCED (raw field), CALCULATED (formula), HYBRID (both)'],
    ['domain_ids *', '', 'One or more: CR, EL, FP, CM, PA, GO, ST, RC (see Section 8)'],
    ['regulatory_references', '', 'E.g., FR Y-14Q, CCAR, DFAST, Basel III, OCC Handbook'],
    ['number_of_instances', '', 'Expected row count (e.g., 410 for facility-level, 100 for counterparty)'],
    ['directly_displayed', '', 'true if shown on CRO dashboard, false if intermediate calculation'],
    ['executable_metric_id', '', 'Link to L3Metric ID (e.g., "C005") — fill in after Section E'],
]
add_table_with_style(doc, len(class_fields), 3, class_fields)

# Section C
doc.add_heading('4.3  Section C: Ingredient Fields (Source Data)', level=2)
doc.add_paragraph(
    'List every atomic field from L1/L2 tables that feeds into this metric\'s calculation. '
    'These are the raw inputs before any computation.'
)
add_note(doc, 'Only list L1 and L2 fields. L3 fields are outputs, not inputs.', 'RULE')

doc.add_paragraph()
ingr_header = [
    ['#', 'Layer', 'Table', 'Field', 'Description', 'Data Type', 'Sample Value'],
    ['1', 'L1 / L2', '', '', '', 'DECIMAL(18,2) / BIGINT / VARCHAR / etc.', ''],
    ['2', '', '', '', '', '', ''],
    ['3', '', '', '', '', '', ''],
    ['4', '', '', '', '', '', ''],
    ['5', '', '', '', '', '', ''],
]
add_table_with_style(doc, len(ingr_header), 7, ingr_header)

# Section D
doc.add_heading('4.4  Section D: Level Definitions (Rollup Logic)', level=2)
doc.add_paragraph(
    'Define how the metric is computed at each rollup level. This is the most critical section — '
    'it determines how the metric aggregates from facility to business segment.'
)

doc.add_heading('Level 1: Facility', level=3)
fac_fields = [
    ['Field', 'Value'],
    ['dashboard_display_name', ''],
    ['in_record', 'true / false'],
    ['sourcing_type', 'Raw / Calc / Agg / Avg'],
    ['level_logic', '(Describe: For each DISTINCT(facility_id) THEN [formula])'],
    ['source_references', '(List layer, table, field, description for each source used at this level)'],
]
add_table_with_style(doc, len(fac_fields), 2, fac_fields)

doc.add_paragraph()
doc.add_heading('Level 2: Counterparty', level=3)
cp_fields = [
    ['Field', 'Value'],
    ['dashboard_display_name', ''],
    ['in_record', 'true / false'],
    ['sourcing_type', 'Raw / Calc / Agg / Avg'],
    ['level_logic', '(Describe aggregation from facility to counterparty)'],
    ['source_references', ''],
]
add_table_with_style(doc, len(cp_fields), 2, cp_fields)

doc.add_paragraph()
doc.add_heading('Level 3: Desk (L3)', level=3)
desk_fields = [
    ['Field', 'Value'],
    ['dashboard_display_name', ''],
    ['in_record', 'true / false'],
    ['sourcing_type', 'Agg / Avg'],
    ['level_logic', '(Describe aggregation from counterparty/facility to desk via enterprise_business_taxonomy)'],
    ['source_references', ''],
]
add_table_with_style(doc, len(desk_fields), 2, desk_fields)

doc.add_paragraph()
doc.add_heading('Level 4: Portfolio (L2)', level=3)
port_fields = [
    ['Field', 'Value'],
    ['dashboard_display_name', ''],
    ['in_record', 'true / false'],
    ['sourcing_type', 'Agg / Avg'],
    ['level_logic', '(Describe aggregation from desk to portfolio via parent_segment_id)'],
    ['source_references', ''],
]
add_table_with_style(doc, len(port_fields), 2, port_fields)

doc.add_paragraph()
doc.add_heading('Level 5: Business Segment (L1)', level=3)
lob_fields = [
    ['Field', 'Value'],
    ['dashboard_display_name', ''],
    ['in_record', 'true / false'],
    ['sourcing_type', 'Agg / Avg'],
    ['level_logic', '(Describe aggregation from portfolio to business segment via parent_segment_id recursion)'],
    ['source_references', ''],
]
add_table_with_style(doc, len(lob_fields), 2, lob_fields)

doc.add_page_break()

# Section E
doc.add_heading('4.5  Section E: Executable Specification (L3Metric)', level=2)
doc.add_paragraph(
    'If this metric needs executable SQL (for the calculation engine), fill in this section. '
    'Skip if the metric is documentation-only or computed externally.'
)

exec_fields = [
    ['Field', 'Value', 'Instructions'],
    ['id *', '', 'Metric ID: C001-C999 (check data/l3-metrics.ts for next available)'],
    ['name *', '', 'Same as item_name above'],
    ['page *', '', 'Dashboard page: P1 (Summary), P2 (Exposure), P3 (Risk), P4 (Performance), P5 (Stress), P6 (Regulatory), P7 (Portfolio)'],
    ['section *', '', 'Section within the page'],
    ['metricType *', '', 'Aggregate, Ratio, Count, Derived, Status, Trend, Table, Categorical'],
    ['formula *', '', 'Human-readable formula (matches generic_formula)'],
    ['displayFormat *', '', '$,.0f (currency), 0.00% (pct), 0,0.00 (decimal), 0,0 (integer)'],
    ['sampleValue *', '', 'Example display value: "$1,234,567" or "1.25x" or "85.3%"'],
]
add_table_with_style(doc, len(exec_fields), 3, exec_fields)

doc.add_paragraph()
doc.add_heading('SQL Formulas by Dimension', level=3)
doc.add_paragraph('Provide SQL for each rollup dimension the metric supports:')

sql_fields = [
    ['Dimension', 'Formula (English)', 'SQL'],
    ['facility', '', '(SELECT ... GROUP BY facility_id)'],
    ['counterparty', '', '(SELECT ... GROUP BY counterparty_id)'],
    ['L3 (Desk)', '', '(SELECT ... JOIN enterprise_business_taxonomy ... GROUP BY segment_id)'],
    ['L2 (Portfolio)', '', '(SELECT ... via parent_segment_id ... GROUP BY segment_id)'],
    ['L1 (LOB)', '', '(SELECT ... recursive parent traversal ... GROUP BY segment_id)'],
]
add_table_with_style(doc, len(sql_fields), 3, sql_fields)

doc.add_paragraph()
doc.add_heading('Source Fields', level=3)
doc.add_paragraph('List every table.field referenced in the SQL above:')
src_fields = [
    ['#', 'Layer', 'Table', 'Field', 'Description'],
    ['1', '', '', '', ''],
    ['2', '', '', '', ''],
    ['3', '', '', '', ''],
    ['4', '', '', '', ''],
]
add_table_with_style(doc, len(src_fields), 5, src_fields)

# Section F
doc.add_heading('4.6  Section F: Demo Data (Optional)', level=2)
doc.add_paragraph(
    'Provide sample walkthrough data if the metric will have an interactive demo in the metric library. '
    'Include 2-3 facility examples with all ingredient field values and the calculated result.'
)

demo_fields = [
    ['Facility', 'Counterparty', 'Input Field 1', 'Input Field 2', '...', 'Calculated Result'],
    ['F-1: [name]', 'CP-1: [name]', '[value]', '[value]', '', '[result]'],
    ['F-2: [name]', 'CP-2: [name]', '[value]', '[value]', '', '[result]'],
    ['F-3: [name]', 'CP-3: [name]', '[value]', '[value]', '', '[result]'],
]
add_table_with_style(doc, len(demo_fields), 6, demo_fields)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 5. SAMPLE FILLED TEMPLATE — DSCR
# ════════════════════════════════════════════════════════════════════════════
doc.add_heading('5. Sample Filled Template — Debt Service Coverage Ratio (DSCR)', level=1)
doc.add_paragraph(
    'Below is a complete worked example showing how to fill out each section '
    'for the Debt Service Coverage Ratio, a key credit risk metric for GSIB banks.'
)

# Section A filled
doc.add_heading('Section A: Metric Identity', level=2)
a_filled = [
    ['Field', 'Value'],
    ['item_id', 'DSCR'],
    ['item_name', 'Debt Service Coverage Ratio'],
    ['abbreviation', 'DSCR'],
    ['kind', 'METRIC'],
    ['definition', 'Net Operating Income divided by Total Debt Service (principal + interest payments due). Measures a borrower\'s ability to service debt obligations from operating cash flow. A DSCR > 1.0x means the borrower generates sufficient income to cover debt payments. Required in FR Y-14Q CRE Schedule and OCC Comptroller\'s Handbook for commercial real estate underwriting.'],
    ['generic_formula', 'Net Operating Income / Total Debt Service'],
    ['insight', 'Measures borrower capacity to meet debt obligations from operating cash flow'],
    ['status', 'ACTIVE'],
]
add_table_with_style(doc, len(a_filled), 2, a_filled)

# Section B filled
doc.add_heading('Section B: Classification & Metadata', level=2)
b_filled = [
    ['Field', 'Value'],
    ['data_type', 'Decimal'],
    ['unit_type', 'RATIO'],
    ['direction', 'HIGHER_BETTER'],
    ['metric_class', 'CALCULATED'],
    ['domain_ids', 'CR, FP'],
    ['regulatory_references', 'FR Y-14Q, CCAR, OCC Comptroller\'s Handbook, CRE underwriting guidelines'],
    ['number_of_instances', '410'],
    ['directly_displayed', 'true'],
    ['executable_metric_id', 'C003'],
]
add_table_with_style(doc, len(b_filled), 2, b_filled)

# Section C filled
doc.add_heading('Section C: Ingredient Fields', level=2)
c_filled = [
    ['#', 'Layer', 'Table', 'Field', 'Description', 'Data Type', 'Sample Value'],
    ['1', 'L2', 'facility_financial_snapshot', 'net_operating_income_amt', 'Annual NOI from the borrower\'s operating statements', 'DECIMAL(18,2)', '12500000.00'],
    ['2', 'L2', 'facility_financial_snapshot', 'total_debt_service_amt', 'Annual P&I payments due on all debt', 'DECIMAL(18,2)', '8750000.00'],
    ['3', 'L1', 'facility_master', 'facility_id', 'Unique facility identifier (join key)', 'BIGINT', '1001'],
    ['4', 'L1', 'facility_master', 'counterparty_id', 'Borrower FK (for counterparty rollup)', 'BIGINT', '42'],
    ['5', 'L2', 'facility_exposure_snapshot', 'drawn_amount', 'Current drawn balance (for EAD weighting)', 'DECIMAL(18,2)', '35000000.00'],
]
add_table_with_style(doc, len(c_filled), 7, c_filled)

# Section D filled
doc.add_heading('Section D: Level Definitions', level=2)

doc.add_heading('Level 1: Facility', level=3)
d_fac = [
    ['Field', 'Value'],
    ['dashboard_display_name', 'Facility DSCR (x)'],
    ['in_record', 'true'],
    ['sourcing_type', 'Calc'],
    ['level_logic', 'For each DISTINCT(facility_id) THEN [net_operating_income_amt] / [total_debt_service_amt]. Result expressed as a ratio (e.g., 1.43x). If total_debt_service = 0, return NULL.'],
    ['source_references', 'L2.facility_financial_snapshot.net_operating_income_amt — Annual NOI\nL2.facility_financial_snapshot.total_debt_service_amt — Annual debt service'],
]
add_table_with_style(doc, len(d_fac), 2, d_fac)

doc.add_paragraph()
doc.add_heading('Level 2: Counterparty', level=3)
d_cp = [
    ['Field', 'Value'],
    ['dashboard_display_name', 'Counterparty DSCR (x)'],
    ['in_record', 'true'],
    ['sourcing_type', 'Avg'],
    ['level_logic', 'For each [counterparty_id] THEN lookup all [facility_id] WHERE IS([counterparty_id]). Weighted average: SUM(facility_DSCR * drawn_amount) / SUM(drawn_amount). Weights by drawn exposure so larger facilities dominate the counterparty-level DSCR.'],
    ['source_references', 'Facility-level DSCR (from Level 1)\nL2.facility_exposure_snapshot.drawn_amount — Weight factor'],
]
add_table_with_style(doc, len(d_cp), 2, d_cp)

doc.add_paragraph()
doc.add_heading('Level 3: Desk (L3)', level=3)
d_desk = [
    ['Field', 'Value'],
    ['dashboard_display_name', 'Desk DSCR (x)'],
    ['in_record', 'true'],
    ['sourcing_type', 'Avg'],
    ['level_logic', 'For each L3 segment in enterprise_business_taxonomy WHERE tree_level=\'L3\': SUM(facility_DSCR * EAD) / SUM(EAD) across all facilities mapped to this desk. Uses facility_segment_mapping to link facilities to organizational units.'],
    ['source_references', 'Facility-level DSCR\nL2.facility_exposure_snapshot.drawn_amount — EAD weight\nL1.enterprise_business_taxonomy — Org hierarchy'],
]
add_table_with_style(doc, len(d_desk), 2, d_desk)

doc.add_paragraph()
doc.add_heading('Level 4: Portfolio (L2)', level=3)
d_port = [
    ['Field', 'Value'],
    ['dashboard_display_name', 'Portfolio DSCR (x)'],
    ['in_record', 'true'],
    ['sourcing_type', 'Avg'],
    ['level_logic', 'For each L2 segment: aggregate child L3 desk DSCRs via parent_segment_id. SUM(desk_DSCR * desk_EAD) / SUM(desk_EAD). Traverses enterprise_business_taxonomy one level up from L3 to L2.'],
    ['source_references', 'Desk-level DSCR\nL1.enterprise_business_taxonomy.parent_segment_id — Hierarchy link'],
]
add_table_with_style(doc, len(d_port), 2, d_port)

doc.add_paragraph()
doc.add_heading('Level 5: Business Segment (L1)', level=3)
d_lob = [
    ['Field', 'Value'],
    ['dashboard_display_name', 'Business Segment DSCR (x)'],
    ['in_record', 'true'],
    ['sourcing_type', 'Avg'],
    ['level_logic', 'For each L1 segment: aggregate child L2 portfolio DSCRs via recursive parent_segment_id traversal. SUM(portfolio_DSCR * portfolio_EAD) / SUM(portfolio_EAD). Top of the rollup hierarchy.'],
    ['source_references', 'Portfolio-level DSCR\nL1.enterprise_business_taxonomy.parent_segment_id — Recursive hierarchy'],
]
add_table_with_style(doc, len(d_lob), 2, d_lob)

doc.add_page_break()

# Section E filled
doc.add_heading('Section E: Executable Specification', level=2)
e_filled = [
    ['Field', 'Value'],
    ['id', 'C003'],
    ['name', 'Debt Service Coverage Ratio'],
    ['page', 'P4 (Financial Performance)'],
    ['section', 'Debt Capacity'],
    ['metricType', 'Ratio'],
    ['formula', 'net_operating_income_amt / total_debt_service_amt'],
    ['displayFormat', '0,0.00x'],
    ['sampleValue', '1.43x'],
]
add_table_with_style(doc, len(e_filled), 2, e_filled)

doc.add_paragraph()
doc.add_heading('SQL — Facility Level', level=3)
add_code_block(doc,
    "SELECT\n"
    "  fm.facility_id,\n"
    "  ffs.net_operating_income_amt / NULLIF(ffs.total_debt_service_amt, 0) AS dscr\n"
    "FROM facility_master fm\n"
    "JOIN facility_financial_snapshot ffs\n"
    "  ON ffs.facility_id = fm.facility_id\n"
    "WHERE ffs.as_of_date = (SELECT MAX(as_of_date) FROM facility_financial_snapshot)\n"
    "GROUP BY fm.facility_id, ffs.net_operating_income_amt, ffs.total_debt_service_amt"
)

doc.add_heading('SQL — Counterparty Level', level=3)
add_code_block(doc,
    "SELECT\n"
    "  fm.counterparty_id,\n"
    "  SUM(ffs.net_operating_income_amt * fes.drawn_amount)\n"
    "    / NULLIF(SUM(ffs.total_debt_service_amt * fes.drawn_amount), 0) AS dscr_weighted\n"
    "FROM facility_master fm\n"
    "JOIN facility_financial_snapshot ffs ON ffs.facility_id = fm.facility_id\n"
    "JOIN facility_exposure_snapshot fes ON fes.facility_id = fm.facility_id\n"
    "WHERE ffs.as_of_date = (SELECT MAX(as_of_date) FROM facility_financial_snapshot)\n"
    "  AND fes.as_of_date = (SELECT MAX(as_of_date) FROM facility_exposure_snapshot)\n"
    "GROUP BY fm.counterparty_id"
)

doc.add_heading('SQL — Desk (L3) Level', level=3)
add_code_block(doc,
    "SELECT\n"
    "  ebt.segment_id,\n"
    "  ebt.segment_name AS desk_name,\n"
    "  SUM(facility_dscr * fes.drawn_amount) / NULLIF(SUM(fes.drawn_amount), 0) AS desk_dscr\n"
    "FROM enterprise_business_taxonomy ebt\n"
    "JOIN facility_segment_mapping fsm ON fsm.segment_id = ebt.segment_id\n"
    "JOIN (\n"
    "  -- facility-level DSCR subquery\n"
    "  SELECT fm.facility_id,\n"
    "    ffs.net_operating_income_amt / NULLIF(ffs.total_debt_service_amt, 0) AS facility_dscr\n"
    "  FROM facility_master fm\n"
    "  JOIN facility_financial_snapshot ffs ON ffs.facility_id = fm.facility_id\n"
    ") sub ON sub.facility_id = fsm.facility_id\n"
    "JOIN facility_exposure_snapshot fes ON fes.facility_id = sub.facility_id\n"
    "WHERE ebt.tree_level = 'L3'\n"
    "GROUP BY ebt.segment_id, ebt.segment_name"
)

# Section F filled
doc.add_heading('Section F: Demo Data', level=2)
demo_filled = [
    ['Facility', 'Counterparty', 'NOI ($)', 'Total Debt Service ($)', 'Drawn ($)', 'DSCR'],
    ['F-1: 125 Park Ave Office', 'CP-1: Meridian Properties', '12,500,000', '8,750,000', '35,000,000', '1.43x'],
    ['F-2: Harbor View Mall', 'CP-1: Meridian Properties', '8,200,000', '9,100,000', '42,000,000', '0.90x'],
    ['F-3: Tech Campus Bldg A', 'CP-2: Apex Developments', '15,800,000', '10,500,000', '55,000,000', '1.50x'],
]
add_table_with_style(doc, len(demo_filled), 6, demo_filled)

doc.add_paragraph()
doc.add_paragraph(
    'Counterparty-level DSCR for CP-1 (weighted by drawn): '
    '(1.43 * 35M + 0.90 * 42M) / (35M + 42M) = (50.05M + 37.80M) / 77M = 1.14x'
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 6. END-TO-END INSTRUCTIONS
# ════════════════════════════════════════════════════════════════════════════
doc.add_heading('6. End-to-End Instructions', level=1)

doc.add_heading('6.1  Step 1: Define the Business Concept', level=2)
doc.add_paragraph('Begin by answering these questions:')
step1 = [
    ['Question', 'Your Answer', 'Guidance'],
    ['What business question does this metric answer?', '', 'E.g., "Can this borrower repay its debt from operating income?"'],
    ['Who consumes this metric?', '', 'CRO, Portfolio Managers, Regulators, Credit Officers'],
    ['What regulatory frameworks reference it?', '', 'FR Y-14Q, CCAR, Basel III, OCC Handbook, DFAST'],
    ['Is higher better or lower better?', '', 'DSCR: higher is better (more income vs. debt). LTV: lower is better.'],
    ['What is the unit?', '', 'Ratio (1.43x), Percentage (85.3%), Currency ($1.2M), Count (42)'],
    ['Is it directly displayed on the CRO dashboard?', '', 'Yes = directly_displayed: true. No = intermediate calculation.'],
]
add_table_with_style(doc, len(step1), 3, step1)

doc.add_heading('6.2  Step 2: Identify Ingredient Fields', level=2)
doc.add_paragraph('Trace the formula back to atomic source fields:')
p = doc.add_paragraph()
p.add_run('1. ').bold = True
p.add_run('Write the formula in plain English (e.g., NOI / Total Debt Service)')
p = doc.add_paragraph()
p.add_run('2. ').bold = True
p.add_run('For each variable in the formula, find which L1 or L2 table contains it')
p = doc.add_paragraph()
p.add_run('3. ').bold = True
p.add_run('Record the exact column name, data type, and a sample value')
p = doc.add_paragraph()
p.add_run('4. ').bold = True
p.add_run('Include join keys (facility_id, counterparty_id) needed to connect tables')
p = doc.add_paragraph()
p.add_run('5. ').bold = True
p.add_run('Include any weight fields needed for rollup (drawn_amount, EAD)')

add_note(doc, 'If a field doesn\'t exist in L1/L2 yet, it must be added to the data dictionary first. Never reference L3 tables as inputs — those are outputs.', 'IMPORTANT')

doc.add_heading('6.3  Step 3: Design Rollup Logic', level=2)
doc.add_paragraph(
    'For each level in the hierarchy, determine how the metric aggregates. '
    'Use the Sourcing Type Decision Tree (Appendix A) to pick the right type.'
)

doc.add_paragraph()
doc.add_heading('Common Rollup Patterns for GSIB Metrics', level=3)
patterns = [
    ['Pattern', 'Metric Types', 'Facility → Counterparty', 'Counterparty → Desk → Portfolio → LOB'],
    ['Simple SUM', 'Exposure, Revenue, Count metrics', 'SUM of facility values', 'SUM up the hierarchy tree'],
    ['Weighted Average', 'DSCR, PD, LGD, Interest Rate', 'SUM(metric * weight) / SUM(weight)', 'Same weighted average, using level-appropriate weight'],
    ['MAX / MIN', 'Days Past Due, Worst Rating', 'MAX across facilities', 'MAX across children'],
    ['COUNT', 'Number of facilities, events', 'COUNT of facilities', 'SUM of counts from children'],
    ['Ratio of SUMs', 'LTV, Utilization Rate', 'Calc: numerator / denominator', 'SUM(numerator) / SUM(denominator) — NOT average of ratios'],
]
add_table_with_style(doc, len(patterns), 4, patterns)

add_note(doc, 'For ratio metrics, always aggregate as SUM(numerator)/SUM(denominator), never as AVG(facility_ratio). Averaging ratios produces incorrect results when facility sizes differ.', 'CRITICAL')

doc.add_heading('6.4  Step 4: Write Executable SQL', level=2)
doc.add_paragraph(
    'Write SQL for each supported dimension. Follow these rules:'
)
rules = [
    '1. Facility SQL is the base — all other levels build on it',
    '2. Counterparty level: JOIN facility_master to get counterparty_id, then GROUP BY counterparty_id',
    '3. Desk (L3): JOIN enterprise_business_taxonomy WHERE tree_level=\'L3\' via facility_segment_mapping',
    '4. Portfolio (L2): Navigate up via parent_segment_id from L3 segments',
    '5. LOB (L1): Recursive parent_segment_id traversal to find L1 root segments',
    '6. Always use NULLIF to prevent division by zero',
    '7. Always filter to the latest as_of_date with a subquery',
    '8. Use table aliases consistently (fm=facility_master, fes=facility_exposure_snapshot, etc.)',
]
for rule in rules:
    doc.add_paragraph(rule, style='List Bullet')

doc.add_heading('6.5  Step 5: Validate & Review', level=2)
doc.add_paragraph('Run through this validation checklist before submitting:')

val_checks = [
    ['#', 'Check', 'Pass?', 'Notes'],
    ['1', 'item_id is unique (not already in catalogue.json)', '', ''],
    ['2', 'All ingredient_fields exist in the data dictionary', '', ''],
    ['3', 'Level definitions cover all 5 rollup levels', '', ''],
    ['4', 'Sourcing types are correct (Calc at facility, Agg/Avg at higher levels)', '', ''],
    ['5', 'SQL compiles without errors against the schema', '', ''],
    ['6', 'SQL produces correct results for sample data', '', ''],
    ['7', 'Division-by-zero is handled (NULLIF)', '', ''],
    ['8', 'Weighted average uses appropriate weights (drawn_amount or EAD)', '', ''],
    ['9', 'Display format matches unit_type ($ for currency, % for percentage, x for ratio)', '', ''],
    ['10', 'Regulatory references are accurate', '', ''],
    ['11', 'Domain assignments are correct (see Section 8)', '', ''],
    ['12', 'Metric tells a coherent story at every rollup level', '', ''],
]
add_table_with_style(doc, len(val_checks), 4, val_checks)

doc.add_heading('6.6  Step 6: Register in the System', level=2)
doc.add_paragraph('After validation, register the metric in these files:')

reg_steps = [
    ['Step', 'File', 'Action'],
    ['1', 'data/metric-library/catalogue.json', 'Add the CatalogueItem JSON object (Sections A-D)'],
    ['2', 'data/l3-metrics.ts', 'Add the L3Metric TypeScript object (Section E) — assign next available C-ID'],
    ['3', 'data/metric-library/catalogue.json', 'Set executable_metric_id to the C-ID from step 2'],
    ['4', 'npm run test:metrics', 'Run metric validation to check for errors'],
    ['5', 'npm run test:calc-engine', 'Run calculation engine tests'],
    ['6', 'app/metrics/[metric]-lineage/ (optional)', 'Add custom lineage visualization page if needed'],
]
add_table_with_style(doc, len(reg_steps), 3, reg_steps)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 7. ROLLUP REVIEW CHECKLIST FOR GSIB
# ════════════════════════════════════════════════════════════════════════════
doc.add_heading('7. Rollup Review Checklist for GSIB', level=1)
doc.add_paragraph(
    'Use this checklist when reviewing a metric\'s rollup definitions for GSIB compliance. '
    'Each item must be verified at every level of the hierarchy.'
)

doc.add_heading('7.1  Mathematical Correctness', level=2)
math_checks = [
    ['Check', 'What to Verify', 'Common Failure'],
    ['Additive metrics sum correctly', 'Facility exposures sum to counterparty, counterparty to desk, etc.', 'Double-counting when facilities belong to multiple segments'],
    ['Ratios aggregate as ratio-of-sums', 'LTV = SUM(committed) / SUM(collateral), not AVG(facility_LTV)', 'Averaging ratios gives wrong answer for unequal-sized facilities'],
    ['Weighted averages use correct weights', 'PD weighted by EAD, DSCR weighted by drawn_amount', 'Using equal weights when facilities have very different sizes'],
    ['Counts are non-overlapping', 'A facility is counted exactly once at each level', 'Facility counted in multiple desks due to mapping error'],
    ['NULL handling is consistent', 'NULL inputs produce NULL output, not 0 or infinity', 'Division by zero returns infinity instead of NULL'],
]
add_table_with_style(doc, len(math_checks), 3, math_checks)

doc.add_heading('7.2  Data Lineage Integrity', level=2)
lineage_checks = [
    ['Check', 'What to Verify'],
    ['Source fields are in L1/L2 only', 'No L3 table appears as an ingredient_field — L3 is output only'],
    ['FK chain is unbroken', 'facility → credit_agreement → counterparty chain has no NULL joins'],
    ['Org hierarchy traversal works', 'enterprise_business_taxonomy correctly resolves L3→L2→L1 via parent_segment_id'],
    ['Date alignment', 'All snapshots use the same as_of_date (or latest available)'],
    ['Schema consistency', 'All referenced tables and columns exist in the data dictionary'],
]
add_table_with_style(doc, len(lineage_checks), 2, lineage_checks)

doc.add_heading('7.3  Regulatory Alignment', level=2)
reg_checks = [
    ['Check', 'What to Verify'],
    ['FR Y-14Q alignment', 'Metric definition matches FR Y-14Q field specifications where applicable'],
    ['CCAR/DFAST consistency', 'Stressed values use the same rollup logic as base values'],
    ['Basel III conformance', 'Risk-weight calculations follow standardized or IRB approach correctly'],
    ['Audit trail', 'Every calculated value can be traced back to atomic L1/L2 inputs'],
    ['Threshold compliance', 'Metric thresholds (from metric_threshold table) are applied at the correct rollup level'],
]
add_table_with_style(doc, len(reg_checks), 2, reg_checks)

doc.add_heading('7.4  GSIB-Specific Considerations', level=2)
gsib_checks = [
    ['Consideration', 'Impact on Rollup'],
    ['Multi-entity structure', 'Counterparties may appear under multiple legal entities — ensure no double-counting'],
    ['Cross-border exposure', 'Currency conversion must happen before aggregation, not after'],
    ['Participation/syndication', 'Bank\'s share (participation_pct) must be applied at facility level, before rollup'],
    ['Off-balance-sheet items', 'Undrawn commitments, guarantees must be included with appropriate credit conversion factors'],
    ['Netting agreements', 'ISDA netting must be applied at the counterparty level, not facility level'],
    ['Consolidated vs. solo', 'Rollup must support both consolidated (group) and solo (entity) views'],
    ['Intraday exposure', 'Settlement risk and intraday exposure may require separate rollup paths'],
]
add_table_with_style(doc, len(gsib_checks), 2, gsib_checks)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 8. DOMAIN REFERENCE
# ════════════════════════════════════════════════════════════════════════════
doc.add_heading('8. Domain Reference', level=1)
doc.add_paragraph('Assign one or more domains to classify your metric:')

domains = [
    ['Domain ID', 'Domain Name', 'Description', 'Example Metrics'],
    ['CR', 'Credit Risk', 'PD, LGD, EAD, expected loss, ratings, delinquency', 'PD, LGD, EAD, DSCR, Days Past Due'],
    ['EL', 'Exposure & Limits', 'Exposure balances, committed amounts, utilization, limits', 'Total Exposure, Utilization Rate, Undrawn Exposure'],
    ['FP', 'Financial Performance', 'Profitability, pricing, return metrics', 'ROE, ROA, NIM, DSCR, Net Interest Income'],
    ['CM', 'Collateral & Mitigation', 'Collateral valuation, LTV, coverage ratios', 'LTV, Collateral Coverage, Haircut %'],
    ['PA', 'Portfolio Analytics', 'Concentration, diversification, portfolio composition', 'HHI, Top-10 Exposure, Sector Concentration'],
    ['GO', 'Governance & Operations', 'Operational metrics, compliance, process KPIs', 'Exception Count, Approval Turnaround'],
    ['ST', 'Stress Testing', 'CCAR/DFAST stressed values, scenario results', 'Stressed PD, Stressed LGD, Capital Impact'],
    ['RC', 'Regulatory Capital', 'RWA, capital ratios, buffers', 'CET1 Ratio, RWA, Capital Buffer'],
]
add_table_with_style(doc, len(domains), 4, domains)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 9. FIELD NAMING CONVENTION
# ════════════════════════════════════════════════════════════════════════════
doc.add_heading('9. Field Naming Convention Reference', level=1)
doc.add_paragraph(
    'Column names in the data model implicitly declare their data type via suffix. '
    'Use these conventions when defining new ingredient fields.'
)

naming = [
    ['Suffix', 'PostgreSQL Type', 'Example Column'],
    ['_id', 'BIGINT', 'counterparty_id, facility_id'],
    ['_code', 'VARCHAR(30)', 'currency_code, pricing_tier_code'],
    ['_name, _desc, _text', 'VARCHAR(500)', 'facility_name, event_desc'],
    ['_amt', 'NUMERIC(20,4)', 'committed_facility_amt, drawn_amt'],
    ['_pct', 'NUMERIC(10,6)', 'coverage_ratio_pct, utilization_pct'],
    ['_value', 'NUMERIC(12,6)', 'collateral_value, market_value'],
    ['_date', 'DATE', 'maturity_date, origination_date'],
    ['_ts', 'TIMESTAMP', 'created_ts, updated_ts'],
    ['_flag', 'BOOLEAN', 'is_active_flag, is_default_flag'],
    ['_count', 'INTEGER', 'number_of_loans, event_count'],
    ['_bps', 'NUMERIC(10,4)', 'interest_rate_spread_bps'],
    ['(fallback)', 'VARCHAR(64)', 'Any column not matching above suffixes'],
]
add_table_with_style(doc, len(naming), 3, naming)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# APPENDIX A: SOURCING TYPE DECISION TREE
# ════════════════════════════════════════════════════════════════════════════
doc.add_heading('Appendix A: Sourcing Type Decision Tree', level=1)
doc.add_paragraph('Use this decision tree to determine the correct sourcing_type for each level definition:')

doc.add_paragraph()
tree_text = (
    "START: Is this the facility (lowest) level?\n"
    "  ├── YES: Is the value read directly from a single L1/L2 field?\n"
    "  │     ├── YES → Raw\n"
    "  │     └── NO  → Calc (formula applied to atomic fields)\n"
    "  └── NO: Is this a higher aggregation level?\n"
    "        ├── YES: Is the metric additive (SUM makes sense)?\n"
    "        │     ├── YES → Agg (simple aggregation: SUM, COUNT, MAX)\n"
    "        │     └── NO:  Is it a ratio or rate?\n"
    "        │           ├── YES → Avg (weighted average using EAD or drawn)\n"
    "        │           └── NO  → Calc (re-calculation from aggregated components)\n"
    "        └── (Should not reach here — all levels are either facility or higher)"
)
add_code_block(doc, tree_text)

doc.add_paragraph()
doc.add_heading('Decision Examples', level=2)
examples = [
    ['Metric', 'Facility Level', 'Counterparty Level', 'Desk/Portfolio/LOB Level'],
    ['Total Exposure ($)', 'Raw (read from facility_exposure_snapshot)', 'Agg (SUM of facility exposures)', 'Agg (SUM up hierarchy)'],
    ['DSCR (x)', 'Calc (NOI / debt_service)', 'Avg (EAD-weighted)', 'Avg (EAD-weighted)'],
    ['LTV (%)', 'Calc (committed / collateral)', 'Calc (SUM committed / SUM collateral)', 'Calc (SUM committed / SUM collateral)'],
    ['Facility Count', 'Raw (always 1)', 'Agg (COUNT of facilities)', 'Agg (SUM of counts)'],
    ['PD (%)', 'Raw (from credit model)', 'Avg (EAD-weighted)', 'Avg (EAD-weighted)'],
    ['Days Past Due', 'Raw (from delinquency snapshot)', 'Agg (MAX across facilities)', 'Agg (MAX across children)'],
]
add_table_with_style(doc, len(examples), 4, examples)

doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('— End of Document —')
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.font.size = Pt(12)

# ── Save ────────────────────────────────────────────────────────────────────
output_path = os.path.expanduser('~/Downloads/Dynamic_Metric_Generation_System_Template.docx')
doc.save(output_path)
print(f'Document saved to: {output_path}')
